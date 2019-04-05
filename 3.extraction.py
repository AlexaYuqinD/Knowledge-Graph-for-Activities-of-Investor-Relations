# -*- coding: utf-8 -*-
"""
Created on Mon Apr  2 14:19:38 2018

@author: THINKPAD
"""

from docx import Document
from win32com.client import Dispatch, constants
import re
import pandas as pd
import os
#import pymysql
#from sqlalchemy import create_engine  

#pymysql.install_as_MySQLdb()
# 处理guests的name和company
def get_guests(spl2):
    guest = {}
    count = 0
    companies = []
    company = ''
    #guests = pd.DataFrame(columns = ['title','name'])
    for mix in spl2:
        if mix.endswith('排序'):
            continue
        if '等' in mix:
            for k in range(1,len(mix)-1):
                tempend1 = mix[-k]
                if tempend1 == '等':
                    mix = mix[:-k]
                    break
        #end = mix[-2:]
        if mix.endswith(('公司','基金','证券','资本','资产','保险','管理','投资','控股','集团','银行','资管','财保','柏瑞','国泰君安','创投','研究','证券报','财富盛世','前海觅贝','深度价值','研究所','中植东方','汇垠天星','分行','工银瑞信','申万宏源','投信','信托','理事会','财险','广证恒生','逸原达','中信建投','国投安信','华夏未来','华泰瑞联','部','第一创业','交银施罗德','安信乾盛','广州金控','国寿安保','前海道谊','租赁','浙江宝基','人寿','易方达','中银国际')):
            company = mix
            companies.append(mix)
            count += 1
        elif len(mix) < 5:
            guest[mix] = company
        else:
            for s in range(2,len(mix)-1):
                tempend = mix[:-s]
                if tempend.endswith(('公司','基金','证券','资本','资产','保险','管理','投资','控股','集团','银行','资管','财保','柏瑞','国泰君安','创投','研究所','分行','工银瑞信','广证恒生','嘉实国际','东方港湾','招商电子','申万宏源','投信','信托','理事会','财险','部','交银施罗德','国寿安保','前海道谊','租赁','浙江宝基','人寿','中信建投','易方达','中银国际')):
                    name = mix[-s:]
                    company = mix[:-s]
                    guest[name] = company
                    break
            if s == len(mix)-2:
                guest[mix] = ' '
                print(mix)
            
    if count == len(spl2):    
        na = [' ' for n in range(count)]
        guests = pd.DataFrame({'name':na,'company':companies},columns = ['name','company'])   
    else:    
        guests = pd.DataFrame.from_dict(guest,orient = 'index')
        guests = guests.reset_index()
        guests.columns = ['name','company']
        
    return guests

# 处理hosts的name和position

def get_hosts_old(spl3):
    pos = ''
    host = {}
    count1 = 0
    position = []
    for mix1 in spl3:
        if mix1.endswith('排序'):
            continue
        if mix1.endswith(('秘书','部长','总监','经理','事长','负责人','部','代表','员','助理','董事','行长','董秘','办公室','工程师','关系','副总','总裁','财务官')):
            pos = mix1
            position.append(mix1)
            count1 += 1
        elif len(mix1) < 5:
            if mix1.endswith(('先生','女士')):
                mix1 = mix1[:-2]
            host[mix1] = pos
        elif len(mix1) == 5 and mix1.endswith(('先生','女士')):
            mix1 = mix1[:-2]
            host[mix1] = pos
        else:
            for s in range(1,len(mix1)-1):
                tempend = mix1[:-s]
                if tempend.endswith(('秘书','部长','总监','经理','事长','负责人','部','代表','员','助理','董事','行长','董秘','办公室','工程师','关系','副总','总裁','财务官')):
                    name1 = mix1[-s:]
                    if name1.endswith(('先生','女士')):
                        name1 = name1[:-2]
                    pos = mix1[:-s]
                   # position.append(pos)
                    host[name1] = pos
                    break
            if s == len(mix1)-2:
                host[mix1] = ''
    
    if count1 == len(spl3):    
        na = [' ' for n in range(count1)]
        hosts = pd.DataFrame({'name':na,'position':position},columns = ['name','position'])   
    else:    
        hosts = pd.DataFrame.from_dict(host,orient = 'index')
        hosts = hosts.reset_index()
        hosts.columns = ['name','position']        
    return hosts

def get_hosts(spl3):
    pos = ''
    host = {}
    count1 = 0
    position = []
    for mix1 in spl3:
        if mix1.endswith('排序'):
            continue
        if mix1.endswith(('秘书','部长','总监','经理','事长','负责人','部','代表','员','助理','董事','行长','董秘','办公室','工程师','关系','副总','总裁','财务官')):
            pos = mix1
            position.append(mix1)
            count1 += 1
            continue
        if mix1.endswith(('先生','女士')):
            mix1 = mix1[:-2]
        if len(mix1) < 5:
            host[mix1] = pos
        else:
            for s in range(1,len(mix1)-1):
                tempend = mix1[:-s]
                if tempend.endswith(('秘书','部长','总监','经理','事长','负责人','部','代表','员','助理','董事','行长','董秘','办公室','工程师','关系','副总','总裁','财务官')):
                    name1 = mix1[-s:]
                    pos = mix1[:-s]
                   # position.append(pos)
                    host[name1] = pos
                    break
            if s == len(mix1)-2:
                host[mix1] = ''
    
    if count1 == len(spl3):    
        na = [' ' for n in range(count1)]
        hosts = pd.DataFrame({'name':na,'position':position},columns = ['name','position'])   
    else:    
        hosts = pd.DataFrame.from_dict(host,orient = 'index')
        hosts = hosts.reset_index()
        hosts.columns = ['name','position']        
    return hosts

def get_eventtype(spl1):
    eventtype = ''
    spl11 = []
    for types in spl1:
        if '□' in types:
            res = re.split('□',types)
            if len(res) == 2:
                spl11.append(res[0])
                spl11.append('□'+res[1])
            else:
                spl11.append(types)
        else:
            spl11.append(types)
            
    while '' in spl11:
        spl11.remove('')       
        
    for types in spl11:
        if len(types) != 0:
            #print(types)
            if types[0] == '（' :
                if len(types)<3:
                    continue
                if types[1] == '请':
                    continue
                if len(eventtype)>0:
                    eventtype = eventtype + ',' + types[1:-1]
                else:
                    eventtype = types[1:-1]
            elif types[0] != '□' and types[1:3] != '其他' and len(types)>2:
                if len(eventtype)>0:
                    if types[0] >= u"\u4e00" and types[0] <= u"\u9fa6":
                        eventtype = eventtype + ',' + types
                    elif types[1] >= u"\u4e00" and types[1] <= u"\u9fa6":
                        eventtype = eventtype + ',' + types[1:]
                else:
                    if types[0] >= u"\u4e00" and types[0] <= u"\u9fa6":
                        eventtype = types
                    elif types[1] >= u"\u4e00" and types[1] <= u"\u9fa6":
                        eventtype = types[1:]
    return eventtype

num = 0
path = os.getcwd()[:-4] + '数据与实验结果\\'
filename = pd.read_excel(path + 'data\\filename.xlsx')[0].tolist()
dateid = pd.read_excel(path + 'data\\date.xlsx')[0].tolist()
#files = os.listdir(path)
#for file in files:
#    if os.os.path.splitext(file)[1] 
f = path+'files\\'+filename[num]
time = dateid[num]
if f.endswith('.DOCX'):
    doc = Document(f)
    t = doc.tables[0]
    p = doc.paragraphs
    newp = []
    info3 = t.cell(0,1).text
    info4 = t.cell(1,1).text
    #time = t.cell(2,1).text
    place = t.cell(3,1).text
    info5 = t.cell(4,1).text
    for ps in p:
        if len(ps.text) != 0:
            newp.append(ps.text)
if f.endswith('.DOC'):
    w = Dispatch('Word.Application')
    doc = w.Documents.Open( FileName = f )
    t = doc.Tables[0]
    p = doc.Paragraphs
    newp = []
    info3 = t.Rows[0].Cells[1].Range.Text
    if info3.endswith('\x07'):
        info3 = info3[:-2]
    info4 = t.Rows[1].Cells[1].Range.Text
    if info4.endswith('\x07'):
        info4 = info4[:-2]
#    time = t.Rows[2].Cells[1].Range.Text
#    if time.endswith('\x07'):
#        time = time[:-2]
    place = t.Rows[3].Cells[1].Range.Text
    if place.endswith('\x07'):
        place = place[:-2]
    info5 = t.Rows[4].Cells[1].Range.Text
    if info5.endswith('\x07'):
        info5 = info5[:-2]
    for ps in p:
        if len(ps.Range.Text) != 0:
            newp.append(ps.Range.Text)
    doc.Close()

while '' in newp:
    newp.remove('')
while '\r' in newp:
    newp.remove('\r')
    
info1 = newp[0]
spl = re.split(r'\s|：',info1)
while '' in spl:
    spl.remove('')
code = spl[1]
simple = spl[3]
#trial = re.search(r'证券简称：(.*)',info1)
#simple = trial.group().split('：')[1]
info2 = newp[1]
#companyname = re.search(r'(.*)公司',info2).group()
companyname = re.search(r'(.*)公司',info2).group()

spl1 = info3.split(' ')
spl1 = re.split(r'\s',info3)

spl2 = re.split(r'-|—|：|:|、|，|,|。|；|/|;|（|）|\s',info4)
while '' in spl2:
    spl2.remove('')
    
spl3 = re.split(r'-|—|：|:|、|，|,|。|；|/|;|（|）|\s',info5)
while '' in spl3:
    spl3.remove('')
eventid = simple + '-' + time
eventtype = get_eventtype(spl1)
event = pd.DataFrame([eventid, companyname,code,eventtype,place],index = ['eventid','companyname','code','eventtype','place']).T 
 
guests = get_guests(spl2)
guests['eventid'] = eventid
hosts = get_hosts(spl3)     
hosts['companycode'] = code
hosts['eventid'] = eventid
event.to_csv(path + 'data\\event.csv',mode = 'a',header = False)
guests.to_csv(path + 'data\\guests.csv',mode = 'a',header = False)
hosts.to_csv(path + 'data\\hosts.csv',mode = 'a',header = False)

#conn = pymysql.Connect(host='127.0.0.1',port=3306,user='root',passwd='qpwoeiruty',db='imooc',charset='utf8')
#cursor = conn.cursor()
#yconnect = create_engine('mysql+mysqldb://root:qpwoeiruty@127.0.0.1:3306/test1?charset=utf8')  
#pd.io.sql.to_sql(event,'event', yconnect, schema='test1', if_exists='append') 
#db = pymysql.connect(host="localhost",user="root",password="qpwoeiruty",db="test1",port=3306, charset = 'utf8')
#cur = db.cursor()
#sql = "select * from event"
#cur.execute(sql)
#results = cur.fetchall()
#db.close()
